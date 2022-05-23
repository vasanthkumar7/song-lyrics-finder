from tkinter import *
from tkinter import messagebox
from docx import Document
import requests
from bs4 import BeautifulSoup
import random



root=Tk()
root.title("Lyrics Finder")
myscroll=Scrollbar(root,orient=VERTICAL)


lables2=[]
numline=0


def new_lyrics():
    global numline
    global lables
    global frame1
    global removebtn
    global count2
    global lables2
    global songbox
    global e1
    rem()
    songname=e1.get().replace(" ","+")
    fgh=[]
    html_file=requests.get("https://search.azlyrics.com/search.php?q="+songname).text
    soup=BeautifulSoup(html_file,"lxml")
    match=[x.extract() for x in soup.findAll('td',class_='text-left visitedlyr')]

    if match==[]:
        messagebox.showinfo("Lyric Finder","song not found ")
        return 
    link=str(list(match[0])[1]).split('"')
    html_files=requests.get(link[1]).text
    soup=BeautifulSoup(html_files,"lxml")
    match=[x.extract() for x in soup.findAll('div',class_='col-xs-12 col-lg-8 text-center')]
    lyrics=str(match).split("\n")
    start=lyrics.index('<!-- Usage of azlyrics.com content by any third-party lyrics provider is prohibited by our licensing agreement. Sorry about that. -->\r')
    end=lyrics.index('<!-- MxM banner -->')
    
    for i in lyrics[start+1:end-2]:
        f=i.replace("<br/>","")
        songbox.insert(END,"          "+f)

    
    lables2=lyrics[start+1:end-2]
    numline=len(lables2)

    removebtn.grid(row=0,column=3,padx=2)
    save.grid(row=0,column=4,padx=2)
    
    
def sa():
    global lables2
    global numline
    global lynum
    global e1
    lynum=random.randint(0, 100)
    na=e1.get()
    name=na+"("+str(lynum)+").docx"
    doc =Document()
    doc.add_heading(na)
    
    with open(name,"a") as f:
        for i in range(numline):
            try:
                encoded=lables2[i].encode('utf-8')
                doc.add_paragraph(lables2[i].replace("<br/>",""))
            except:
                pass
    doc.save(name)
    messagebox.showinfo("lyric finder","song lyrics saved as "+name)
    

def rem():
    global numline
    global lables
    global removebtn
    global lables2
    global songbox
    global e1
    lables2=[]
    songbox.delete(0,numline-1)
    
    removebtn.grid_forget()
    save.grid_forget()
    
    
    
frame1=Frame(root)
labe=Label(frame1,text="Enter the song name:")
labe.grid(row=0,column=0)
e1=Entry(frame1,width=30)
e1.grid(row=0,column=1,padx=5)


sea=Button(frame1,text="Search",command=new_lyrics)
sea.grid(row=0,column=2,padx=2)


frame2=Frame(root)
myscroll2=Scrollbar(frame2,orient=HORIZONTAL)
myscroll=Scrollbar(frame2,orient=VERTICAL)
songbox=Listbox(frame2,width=70,height=30,borderwidth=0,bg="#fce303",fg="black",yscrollcommand=myscroll.set,font=("arial",13),xscrollcommand=myscroll2.set)
myscroll.config(command=songbox.yview)
myscroll2.config(command=songbox.xview)
myscroll.pack(side=RIGHT,fill=Y)
myscroll2.pack(side=BOTTOM,fill=X)
songbox.pack()

frame2.grid(row=1,column=0,pady=10)

removebtn=Button(frame1,text="Clear",command=rem)
save=Button(frame1,text="Save",command=sa)
frame1.grid(row=0,column=0)
root.mainloop()
